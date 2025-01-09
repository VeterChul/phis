import numpy as np
import json
import os
import shutil
from conf import prog_dir
from random import *
from sklearn.utils import shuffle
import docx 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def print_task(js):
    n = len(js["tasks_ans"])
    m = len(js["tasks_ans"][0])
    print("Работа:")
    for i in range(n):
        print(f"Вариант N{i+1}")
        for j in range(m):
            print(f"    Задача N{j + 1}\n    {js["tasks_ans"][i][j]["text"]}")
    print("Ответы:")
    for i in range(n):
        print(f"Вариант N{i+1}")
        for j in range(m):
            print(f"    Задача N{j + 1}\n    {js["tasks_ans"][i][j]["ans"]}")
    

def name_job(prog_dir, js):
    if "job_name" in js:
        if js["job_name"] in os.listdir(prog_dir):
            shutil.rmtree(f"{prog_dir}{js["job_name"]}")
        return js["job_name"]
    else:
        l = os.listdir(prog_dir)
        max = 0
        for i in l:
            try:
                if int(i) > max:
                    max = int(i)
            except:
                pass
        return str(max + 1)

# считывание сида и листа задач

list_dir = os.listdir(prog_dir)

if "sid.json" in list_dir:
    f  = open(f"{prog_dir}sid.json", "r")
    sid  = json.load(f)
    f.close()
else:
    print("Не обнаружен файл sid. Проверьте файлы")
    exit()
list_dir = os.listdir(prog_dir)

if "sid.json" in list_dir:
    f  = open(f"{prog_dir}sid.json", "r")
    sid  = json.load(f)
    f.close()
else:
    print("Не обнаруружен файл sid")
    exit()

if "list_of_tasks.json" in list_dir:
    f  = open(f"{prog_dir}list_of_tasks.json", "r")
    list_of_tasks  = json.load(f)
    f.close()
else:
    print("Не обнаружен файл list_of_tasks. Проверьте файлы")
    exit()

if "list_of_tasks_user.json" in list_dir:
    f  = open(f"{prog_dir}list_of_tasks_user.json", "r")
    list_of_tasks_user  = json.load(f)
    f.close()
else:
    list_of_tasks_user = {}


#генерация 

list_ret = {
    "tasks_ans":[]
}

if "quantity" in sid:
    n = sid["quantity"]
else:
    n = 1


for k in range(n):
    m = []
    for i in sid["list"]:
        if i["name"] in list_of_tasks:
            task = list_of_tasks[i["name"]]
        elif i["name"] in list_of_tasks_user:
            task = list_of_tasks_user[i["name"]]
        else:
            print(f"Указано неправильное имя задачи {i["name"]}")
            exit()
        
        if "quantity" in i:
            r = int(i["quantity"])
        else:
            r = 1

        for j in range(r):
            exec(task["kod"])
            if "values" in i:
                values = i["values"]
            else:
                values = task["standard_values"]
            s = f"an = {i["name"]}('{i["looking"]}',{values}, {task})"
            exec(s)
            m.append({
                "text" : an[0],
                "ans" : an[1]
            })
    list_ret["tasks_ans"].append(m)

if "random" in sid:
    if sid["random"]:
        list_ret["tasks_ans"] =  shuffle(list_ret["tasks_ans"])

print_task(list_ret)

#запись json

name = name_job(prog_dir, sid)

os.mkdir(f"{prog_dir}{name}")
f = open(f"{prog_dir}{name}/job.json", "a")
f.write(str(list_ret).replace("'", '"'))
f.close()

#запись docx

n = len(list_ret["tasks_ans"])
m = len(list_ret["tasks_ans"][0])
for i in range(n):
    doc = docx.Document()
    title = doc.add_heading(f'Работа: {name}. Вариант N{i+1}.', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for j in range(m):
        para = doc.add_paragraph()
        bold_para = para.add_run(f"Задача N{j + 1}")
        bold_para.bold = True
        doc.add_paragraph(f"{list_ret["tasks_ans"][i][j]["text"]}")
    doc.save(f'{prog_dir}{name}/Работа: {name}. Вариант N{i+1}..docx')

for i in range(n):
    doc = docx.Document()
    title = doc.add_heading(f'Работа: {name}. Вариант N{i+1}. Ответы.', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for j in range(m):
        para = doc.add_paragraph()
        bold_para = para.add_run(f"Задача N{j + 1}")
        bold_para.bold = True
        doc.add_paragraph(f"{list_ret["tasks_ans"][i][j]["ans"]}")
    doc.save(f'{prog_dir}{name}/Работа: {name}. Вариант N{i+1}. Ответы..docx')
